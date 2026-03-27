"""
Dynamic CV compiler — WeasyPrint + Jinja2.

For each application, generates a unique PDF CV with tailored bullets
slotted into the template by source_role key.

source_role keys map directly to experience sections in cv_template.html:
  CISI     → CISI Marketing Intern section
  Todlr    → Todlr Brand Development Associate section
  Evolve   → Evolve Customer Experience Associate section
  Deloitte → Deloitte Audit Intern section
  Airtel   → Airtel Networks Sales & Marketing Intern section

Fallback: if WeasyPrint fails for any reason, log the error and return
the path to the base CV PDF. Never raise — a compile failure must not
block the submission stage.
"""

import logging
import os
import re
import shutil
from pathlib import Path

from jinja2 import Environment, FileSystemLoader, select_autoescape

from generation.writer import GenerationResult, CVBullet

log = logging.getLogger(__name__)


class CVCompiler:

    def __init__(
        self,
        template_path: str | None = None,
        base_cv_path:  str | None = None,
        output_dir:    str | None = None,
    ):
        self.template_path = Path(template_path or os.getenv('CV_TEMPLATE_PATH', './assets/cv_template.html'))
        self.base_cv_path  = Path(base_cv_path  or os.getenv('CV_BASE_PDF_PATH',  './assets/Omokolade_Sobande_CV.pdf'))
        _root              = Path(output_dir    or os.getenv('OUTPUT_DIR',         './output'))
        self.output_dir    = _root / 'cvs'
        self.cl_dir        = _root / 'cover_letters'
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.cl_dir.mkdir(parents=True, exist_ok=True)
        # C2: validate base CV exists on startup
        if not self.base_cv_path.exists():
            log.error(
                'C2: base CV not found at %s — CV upload will fall back to empty path. '
                'Set CV_BASE_PDF_PATH in .env to a valid PDF.',
                self.base_cv_path,
            )

    def compile(
        self,
        job_id: int,
        company: str,
        role: str,
        result: GenerationResult,
    ) -> str:
        """
        Render the Jinja2 template with tailored bullets and convert to PDF.

        Returns the absolute path to the generated PDF.
        On WeasyPrint failure, returns the base CV path as fallback.
        """
        safe_company = _safe_filename(company)
        safe_role    = _safe_filename(role)
        filename     = f'{safe_company}_{safe_role}_CV.pdf'
        output_path  = self.output_dir / filename

        # Always save cover letter to disk alongside the CV
        self._save_cover_letter(safe_company, safe_role, result.cover_letter)

        try:
            html = self._render_template(result)
            self._html_to_pdf(html, output_path)
            log.info('CV compiled: %s', output_path)
            return str(output_path)

        except Exception as exc:
            log.error(
                'WeasyPrint failed for %s/%s: %s — falling back to base CV',
                company, role, exc
            )
            # Fallback: copy base CV to output dir with a distinct name so
            # the pipeline has a concrete file path to upload
            fallback_path = self.output_dir / f'{job_id}_{safe_company}_BASE_CV.pdf'
            try:
                shutil.copy2(self.base_cv_path, fallback_path)
                return str(fallback_path)
            except Exception as copy_exc:
                log.error('Base CV copy also failed: %s — returning base path directly', copy_exc)
                return str(self.base_cv_path)

    def _save_cover_letter(self, safe_company: str, safe_role: str, cover_letter: str) -> None:
        """Write cover letter to output/cover_letters/{Company}_{Role}_cover_letter.txt."""
        cl_path = self.cl_dir / f'{safe_company}_{safe_role}_cover_letter.txt'
        try:
            cl_path.write_text(cover_letter, encoding='utf-8')
            log.info('Cover letter saved: %s', cl_path)
        except Exception as exc:
            log.error('Failed to save cover letter to %s: %s', cl_path, exc)

    def _render_template(self, result: GenerationResult) -> str:
        """Render cv_template.html with tailored bullets injected by source_role."""
        if not self.template_path.exists():
            raise FileNotFoundError(f'CV template not found: {self.template_path}')

        # Build a dict of source_role → list of tailored bullet strings
        bullets_by_role: dict[str, list[str]] = {}
        for b in result.cv_bullets:
            bullets_by_role.setdefault(b.source_role, []).append(b.tailored)

        env = Environment(
            loader=FileSystemLoader(str(self.template_path.parent)),
            autoescape=select_autoescape(['html']),
        )
        template = env.get_template(self.template_path.name)

        return template.render(
            bullets=bullets_by_role,
            cover_letter=result.cover_letter,
        )

    def compile_docx(
        self,
        job_id: int,
        company: str,
        role: str,
        result: GenerationResult,
    ) -> str:
        """
        Generate a DOCX CV with tailored bullets and cover letter.
        Returns the absolute path to the .docx file.
        Raises on failure — caller should handle and fall back to PDF.
        """
        from docx import Document
        from docx.shared import Pt

        safe_company = _safe_filename(company)
        safe_role    = _safe_filename(role)
        filename     = f'{safe_company}_{safe_role}_CV.docx'
        output_path  = self.output_dir / filename

        doc = Document()

        # Cover letter section
        heading = doc.add_heading('Cover Letter', level=1)
        heading.runs[0].font.size = Pt(14)
        for para_text in result.cover_letter.split('\n\n'):
            para_text = para_text.strip()
            if para_text:
                doc.add_paragraph(para_text)

        # CV bullets section
        doc.add_heading('Tailored Experience Highlights', level=1)
        bullets_by_role: dict[str, list[str]] = {}
        for b in result.cv_bullets:
            bullets_by_role.setdefault(b.source_role, []).append(b.tailored)

        for role_key, bullets in bullets_by_role.items():
            doc.add_heading(role_key, level=2)
            for bullet in bullets:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(bullet)

        doc.save(str(output_path))
        log.info('DOCX compiled: %s', output_path)

        # Always save cover letter to disk — consistent with PDF workflow
        self._save_cover_letter(safe_company, safe_role, result.cover_letter)

        return str(output_path)

    def _html_to_pdf(self, html: str, output_path: Path) -> None:
        """Convert rendered HTML to PDF using WeasyPrint."""
        from weasyprint import HTML
        HTML(string=html, base_url=str(self.template_path.parent)).write_pdf(str(output_path))


def _safe_filename(s: str, max_len: int = 30) -> str:
    """Strip non-alphanumeric chars and truncate for use in filenames."""
    safe = re.sub(r'[^\w\s-]', '', s).strip()
    safe = re.sub(r'[\s]+', '_', safe)
    return safe[:max_len]
